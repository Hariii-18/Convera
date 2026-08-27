import io

from docx import Document
from docx.shared import Pt, RGBColor

from app.services.export.conversation_content import ConversationExportDocument

_BRAND_COLOR = RGBColor(0x6D, 0x28, 0xD9)
_META_COLOR = RGBColor(0x66, 0x66, 0x66)


class ConversationDocxExporter:
    content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    file_extension = "docx"

    def render(self, document: ConversationExportDocument) -> bytes:
        doc = Document()

        brand_paragraph = doc.add_paragraph()
        brand_run = brand_paragraph.add_run(document.brand.upper())
        brand_run.bold = True
        brand_run.font.size = Pt(10)
        brand_run.font.color.rgb = _BRAND_COLOR

        doc.add_heading(document.meeting_title, level=0)

        meta_parts = [document.date_time_ist, document.duration_label, document.participants_label]
        meta_line = " · ".join(part for part in meta_parts if part)
        if meta_line:
            meta_paragraph = doc.add_paragraph(meta_line)
            meta_run = meta_paragraph.runs[0]
            meta_run.font.size = Pt(9)
            meta_run.font.color.rgb = _META_COLOR

        disclaimer_paragraph = doc.add_paragraph()
        disclaimer_run = disclaimer_paragraph.add_run(document.disclaimer)
        disclaimer_run.italic = True
        disclaimer_run.font.size = Pt(8)
        disclaimer_run.font.color.rgb = _META_COLOR

        for turn in document.turns:
            if turn.speaker_label:
                speaker_paragraph = doc.add_paragraph()
                speaker_run = speaker_paragraph.add_run(f"{turn.speaker_label}:")
                speaker_run.bold = True
                speaker_run.font.color.rgb = _BRAND_COLOR
            for text in turn.texts:
                doc.add_paragraph(text)

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
